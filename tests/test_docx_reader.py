"""Tests for the Word (.docx) reader (Dimensions / Phase 6).

Builds a real .docx with python-docx in a temp dir, then asserts the reader extracts
paragraphs and table cells verbatim. No network, no LLM — pure local round-trip.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from core.docx_reader import DocxReadError, read_docx


def _make_docx(path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("Curriculum Vitae — Jane Doe")
    document.add_paragraph("Experience: 5 years in robotics software")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Level"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Expert"
    document.save(str(path))


def test_read_docx_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    doc_path = tmp_path / "cv.docx"
    _make_docx(doc_path)

    result = read_docx(doc_path)

    assert result.doc_type == "docx"
    assert result.extraction_method == "python-docx"
    assert "Jane Doe" in result.text
    assert "5 years in robotics software" in result.text
    # Table rows are tab-separated, verbatim.
    assert "Skill\tLevel" in result.text
    assert "Python\tExpert" in result.text


def test_read_docx_accepts_str_path(tmp_path: Path) -> None:
    doc_path = tmp_path / "memo.docx"
    _make_docx(doc_path)
    result = read_docx(str(doc_path))
    assert result.source_path == doc_path


def test_read_docx_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        read_docx(tmp_path / "does_not_exist.docx")


def test_read_docx_wrong_suffix_raises(tmp_path: Path) -> None:
    not_docx = tmp_path / "note.txt"
    not_docx.write_text("plain text", encoding="utf-8")
    with pytest.raises(DocxReadError):
        read_docx(not_docx)
