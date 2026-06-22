"""Word (.docx) reader (Dimensions / Phase 6): extract text + tables from Word documents.

Part of the Analyst (Recruiting) and Builder (Finance) dimensions: CVs, memos, and internal
reports often arrive as Word files. This is the *reading* counterpart to the existing
PDF/Excel readers — extraction only, no synthesis (that stays in the pipeline).

Uses **python-docx** (already a project dependency, MIT-licensed, fully local/offline). For
high-fidelity layout/table reconstruction on complex documents, the optional Docling adapter
(:mod:`core.docling_reader`) is preferred; this reader is the lightweight, always-available
path that needs no extra install. Mirrors :mod:`core.pdf_reader` (small seams, lazy imports,
fail-fast errors with an exact fix hint).
"""

from __future__ import annotations

from pathlib import Path

from models.research import DocContent


class DocxReadError(Exception):
    """A .docx document could not be read by python-docx (fail fast, SPEC REQ-5)."""


# --------------------------------------------------------------- backend seam
def _extract_docx(path: Path) -> tuple[str, int]:
    """Return (joined text incl. tables, number of non-empty blocks) via python-docx.

    Paragraphs are emitted in document order, then each table is emitted row by row with
    tab-separated cells — so every word and every cell value is preserved verbatim.
    """
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - dependency is declared in requirements
        raise DocxReadError(
            "python-docx is required to read Word (.docx) files but is not installed.\n"
            "Fix: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    return "\n".join(parts).strip(), len(parts)


# --------------------------------------------------------------------- public API
def read_docx(path: Path | str) -> DocContent:
    """Read a Word ``.docx`` file into :class:`DocContent` (paragraphs + tables, in order).

    Args:
        path: Path to a ``.docx`` file.

    Returns:
        :class:`DocContent` with ``doc_type="docx"`` and ``extraction_method="python-docx"``.

    Raises:
        FileNotFoundError: If the path does not exist.
        DocxReadError: If the type is unsupported or the file is not a readable .docx.
    """
    file_path = Path(path)
    if not file_path.is_file():
        raise FileNotFoundError(
            f"Document not found: {file_path.resolve()}\n"
            "Fix: check the path; it must point to an existing .docx file."
        )

    if file_path.suffix.lower() != ".docx":
        raise DocxReadError(
            f"Unsupported document type '{file_path.suffix}' for {file_path.name}. "
            "read_docx handles .docx only — convert legacy .doc to .docx first."
        )

    try:
        text, block_count = _extract_docx(file_path)
    except DocxReadError:
        raise
    except Exception as exc:  # python-docx raises varied errors for corrupt/locked files
        raise DocxReadError(
            f"Could not read {file_path.name} as a Word document: {exc}\n"
            "Fix: ensure it is a valid, non-password-protected .docx file."
        ) from exc

    return DocContent(
        source_path=file_path,
        doc_type="docx",
        text=text,
        page_count=None,  # Word has no fixed page count without rendering the document
        extraction_method="python-docx",
    )
